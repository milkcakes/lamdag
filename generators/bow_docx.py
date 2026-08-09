"""DOCX export for the Budget of Work (BoW): week-by-week curriculum mapping.

Portrait, 8.5in x 13in (Philippine long bond), DepEd letterhead, then a table
of Week / Learning Competency / Code with the Week cell merged per group.
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from generators.watermark import add_docx_watermark


def _set_cell_shading(cell, color):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def _set_cell_border(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:color"), "000000")
        element.set(qn("w:space"), "0")
        tcBorders.append(element)
    tcPr.append(tcBorders)


def _set_cell_width(cell, width):
    cell.width = width
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = tcPr.find(qn("w:tcW"))
    if tcW is None:
        tcW = OxmlElement("w:tcW")
        tcPr.append(tcW)
    tcW.set(qn("w:w"), str(int(width.twips)))
    tcW.set(qn("w:type"), "dxa")


def _para(cell, first=True, align=None):
    if first and cell.paragraphs and not cell.paragraphs[0].runs:
        p = cell.paragraphs[0]
    else:
        p = cell.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = 1.0
    return p


def _add_runs(p, text, size=10, bold=False, italic=False):
    lines = (text or "").split("\n")
    for i, line in enumerate(lines):
        if i > 0:
            run = p.add_run()
            run.add_break()
        run = p.add_run(line)
        run.font.size = Pt(size)
        run.font.name = "Calibri"
        run.bold = bold
        run.italic = italic
    return p


def _add_letterhead(doc, data):
    region = data.get("region") or "Region VII - Central Visayas"
    division = data.get("letterhead_division") or ""
    school = data.get("letterhead_school") or ""

    spec = [("Republic of the Philippines", 12, True, False),
            ("Department of Education", 13, True, False),
            (region, 10.5, False, True)]
    if division:
        spec.append((division, 11, True, False))
    if school:
        spec.append((school, 11, True, False))

    for text, size, bold, italic in spec:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        _add_runs(p, text, size=size, bold=bold, italic=italic)


def _add_rule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _info_table(doc, data):
    """2-column label/value rows under the title (Learning Area, Grade, etc.)."""
    terms = data.get("terms")
    term_text = ", ".join(f"Term {t}" for t in terms) if terms else ""
    rows = [
        ("Learning Area:", data.get("subject", "")),
        ("Grade Level:", data.get("grade_level", "")),
        ("School Year:", data.get("school_year", "")),
        ("Total Competencies:", str(data.get("total", 0))),
    ]
    if term_text:
        rows.insert(3, ("Terms:", term_text))
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    label_w = Inches(1.9)
    value_w = Inches(4.75)
    for label, value in rows:
        row = table.add_row()
        row.cells[0].text = ""
        row.cells[1].text = ""
        for cell in row.cells:
            _set_cell_width(cell, label_w if cell == row.cells[0] else value_w)
        p = _para(row.cells[0])
        _add_runs(p, label, size=10, bold=True)
        p2 = _para(row.cells[1])
        _add_runs(p2, value, size=10)
    return table


def generate_bow_docx(data, output_path, watermark=False):
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.line_spacing = 1.0

    # Portrait, 8.5in x 13in (Philippine long bond), margins per official template
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(13)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(0.8125)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    _add_letterhead(doc, data)
    _add_rule(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(2)
    _add_runs(title, "BUDGET OF WORK", size=16, bold=True)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(8)
    _add_runs(sub, f"SY {data.get('school_year', '')}".strip() or "School Year", size=11, bold=True)

    rows = data.get("rows", [])
    terms = sorted({r.get("term") for r in rows if r.get("term") is not None})
    data["terms"] = terms

    _info_table(doc, data)

    doc.add_paragraph()
    heading = doc.add_paragraph()
    heading.paragraph_format.space_after = Pt(2)
    _add_runs(heading, "Curriculum Mapping by Term and Week", size=12, bold=True)

    table = doc.add_table(rows=0, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = (Inches(0.9), Inches(0.9), Inches(3.9), Inches(0.99))

    def _add_row(texts, bold=False, shade=None):
        row = table.add_row()
        for i, text in enumerate(texts):
            cell = row.cells[i]
            _set_cell_border(cell)
            _set_cell_width(cell, widths[i])
            if shade:
                _set_cell_shading(cell, shade)
            if text:
                p = _para(cell)
                _add_runs(p, text, size=10, bold=bold)
        return row

    _add_row(["Term", "Week", "Learning Competency", "Code"], bold=True, shade="BFBFBF")

    prev_term = prev_key = None
    term_start = week_start = None
    for r in rows:
        term = r.get("term")
        week = r.get("week")
        key = (term, week)

        if term != prev_term:
            # close previous week group, then emit a term header row
            tr = _add_row([], shade="D9D9D9")
            cell = tr.cells[0].merge(tr.cells[1]).merge(tr.cells[2]).merge(tr.cells[3])
            p = _para(cell)
            _add_runs(p, f"TERM {term}", size=10, bold=True)
            term_start = None
            prev_term = term
            prev_key = None

        is_first_week = key != prev_key
        r_ = _add_row([
            "",  # Term is shown by the TERM header rows
            (f"Week {week}" if week is not None and is_first_week else ""),
            r.get("description", ""),
            r.get("code", ""),
        ])
        if is_first_week:
            week_start = r_
            prev_key = key
        elif week_start is not None:
            week_start.cells[1].merge(r_.cells[1])
            prev_key = key

    if watermark:
        add_docx_watermark(doc)

    doc.save(output_path)
    return output_path
