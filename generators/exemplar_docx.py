"""DOCX export for the official DepEd Strengthened Senior High School Lesson
Exemplar (LE): portrait, 8.5in x 13in long bond paper, sections I-VIII with a
two-column PROCEDURES | ANNOTATIONS table."""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from generators import exemplar_layout as L

LABEL_W = Inches(2.15)
VALUE_W = Inches(4.5)
SUM_W = Inches(2.15 + 4.5)
GRAY = "BFBFBF"
LIGHT = "D9D9D9"
PHASE = "DDEBF7"


def _set_cell_shading(cell, color):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def _set_cell_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:color"), "000000")
        element.set(qn("w:space"), "0")
        tcBorders.append(element)
    tcPr.append(tcBorders)


def _set_cell_width(cell, width):
    cell.width = width
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = tcPr.find(qn("w:tcW"))
    if tcW is None:
        tcW = OxmlElement("w:tcW")
        tcPr.append(tcW)
    tcW.set(qn("w:w"), str(int(width.twips)))
    tcW.set(qn("w:type"), "dxa")


def _fit_grid(table, width_a, width_b):
    """Set the table grid so Word honors the intended column widths.

    python-docx creates a tblGrid with equal columns; with fixed layout Word
    uses those grid columns instead of the per-cell w:tcW, so without this the
    columns render 50/50 and the layout breaks. We fix gridCol + tblW here.
    """
    table.columns[0].width = width_a
    table.columns[1].width = width_b
    tblPr = table._tbl.tblPr
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), str(int((width_a + width_b) / 635)))
    tblW.set(qn("w:type"), "dxa")


def _add_spacer(doc, size=4):
    """Minimal empty paragraph between two tables (Word requires a paragraph
    between adjacent tables or it may merge them on open)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run("")
    run.font.size = Pt(size)
    return p


def _para(cell, first=True):
    if first and cell.paragraphs and not cell.paragraphs[0].runs:
        p = cell.paragraphs[0]
    else:
        p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = 1.0
    return p


def _add_runs(p, text, size=10, bold=False, italic=False, underline=False):
    lines = (text or "").split("\n")
    for i, line in enumerate(lines):
        if i > 0:
            run = p.add_run()
            run.add_break()
        run = p.add_run(line)
        run.font.size = Pt(size)
        run.font.name = "Calibri"
        run.bold = bold
        run.italic = italic
        run.underline = underline
    return p


def _value_cell(cell, value, placeholder="\u2014"):
    p = _para(cell)
    _add_runs(p, value if value else placeholder, size=10)


def _label_cell(cell, title, info_style=False):
    p = _para(cell)
    _add_runs(p, title, size=10, bold=True, italic=info_style)


def _add_row(table, width_a=LABEL_W, width_b=VALUE_W):
    row = table.add_row()
    for cell in row.cells:
        _set_cell_border(cell)
    _set_cell_width(row.cells[0], width_a)
    _set_cell_width(row.cells[1], width_b)
    return row


def _section_header(table, text):
    """Gray full-width row for a numbered section header."""
    row = _add_row(table, SUM_W, SUM_W)
    cell = row.cells[0].merge(row.cells[1])
    _set_cell_shading(cell, GRAY)
    p = _para(cell)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_runs(p, text, size=10, bold=True)


def _field_row(table, title, value, info_style=False):
    row = _add_row(table)
    _label_cell(row.cells[0], title, info_style=info_style)
    _value_cell(row.cells[1], value)


def _merged_value(table, text, italic=False, placeholder="\u2014"):
    row = _add_row(table, SUM_W, SUM_W)
    cell = row.cells[0].merge(row.cells[1])
    p = _para(cell)
    _add_runs(p, text if text else placeholder, size=10, italic=italic)


def _objectives_text(data):
    items = data.get("le_objectives") or []
    intro = data.get("le_objectives_intro") or L.OBJECTIVES_INTRO
    if not items:
        return intro
    return intro + "\n" + "\n".join(
        f"{i}. {line}" for i, line in enumerate(items, 1)
    )


def _add_letterhead(doc, data):
    region = data.get("region") or "Region VII - Central Visayas"
    division = data.get("letterhead_division") or ""
    school = data.get("letterhead_school") or ""

    spec = [("Republic of the Philippines", 12, True, False),
            ("Department of Education", 13, True, False),
            (region, 10.5, False, True)]
    if division:
        spec.append((division, 11, True, False))
    if school:
        spec.append((school, 11, True, False))

    for text, size, bold, italic in spec:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        _add_runs(p, text, size=size, bold=bold, italic=italic)


def _add_rule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_exemplar_title(doc, data):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    _add_runs(p, "LESSON EXEMPLAR", size=14, bold=True)

    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tp.paragraph_format.space_before = Pt(0)
    tp.paragraph_format.space_after = Pt(0)
    _add_runs(tp, f"Lesson Title/Topic: {data.get('lesson_name', '') or ''}", size=11, bold=True)

    if data.get("school_year"):
        sp = doc.add_paragraph()
        sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sp.paragraph_format.space_before = Pt(0)
        sp.paragraph_format.space_after = Pt(8)
        _add_runs(sp, f"School Year {data['school_year']}", size=11, italic=True)
    else:
        doc.paragraphs[-1].paragraph_format.space_after = Pt(8)


def _add_procedures(table, data):
    # PROCEDURES | ANNOTATIONS column header (official two-column header row;
    # content column is wide on the left, annotation column narrow on the right)
    row = _add_row(table, VALUE_W, LABEL_W)
    _set_cell_shading(row.cells[0], GRAY)
    _set_cell_shading(row.cells[1], GRAY)
    p0 = _para(row.cells[0])
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_runs(p0, L.HEADER_PROCEDURES, size=10, bold=True)
    p1 = _para(row.cells[1])
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_runs(p1, L.HEADER_ANN_COL, size=10, bold=True)

    for phase in data.get("le_phases", []):
        for i, step in enumerate(phase["steps"], 1):
            srow = _add_row(table, VALUE_W, LABEL_W)
            cp = _para(srow.cells[0])
            _add_runs(cp, f"{i}. {step['title']}", size=10, bold=True)
            vp = _para(srow.cells[0], first=False)
            _add_runs(vp, step["content"] or "\u2014", size=10)
            ap = _para(srow.cells[1])
            if i == 1:
                # Phase label (A./B./C.) goes in the ANNOTATIONS column, as in
                # the official template.
                _add_runs(ap, phase["name"], size=9, bold=True, italic=True)
                ap = _para(srow.cells[1], first=False)
            _add_runs(ap, step["annotation"] or "", size=9, italic=True)


def _add_signature(doc, data):
    doc.add_paragraph()
    sig = doc.add_table(rows=3, cols=2)
    sig.alignment = WD_TABLE_ALIGNMENT.CENTER

    designation = data.get("designation", "") or "Teacher"
    head_designation = data.get("school_head_designation", "") or "School Head / Principal"
    teacher_name = data.get("teacher", "") or "\u00a0"
    head_name = data.get("school_head", "") or "\u00a0"

    sig_rows = [
        (("Prepared by:", True, False), ("Reviewed & Checked by:", True, False)),
        ((teacher_name, False, True), (head_name, False, True)),
        ((designation, False, False), (head_designation, False, False)),
    ]
    for r, cells in enumerate(sig_rows):
        for c, (text, bold, underline) in enumerate(cells):
            cell = sig.rows[r].cells[c]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if r == 1:
                p.paragraph_format.space_before = Pt(6)
            run = p.add_run(text)
            run.font.size = Pt(10)
            run.font.name = "Calibri"
            run.bold = bold
            run.underline = underline


def generate_exemplar_docx(data, output_path):
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.line_spacing = 1.0

    # Portrait, 8.5in x 13in (Philippine long bond), margins per official template
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(13)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(0.8125)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    _add_letterhead(doc, data)
    _add_rule(doc)
    _add_exemplar_title(doc, data)

    def _new_table():
        t = doc.add_table(rows=0, cols=2)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        t.autofit = False
        return t

    # --- Main table: lesson details + I-V ---
    table = _new_table()
    _fit_grid(table, LABEL_W, VALUE_W)

    # --- Lesson details ---
    for label, key in L.INFO_ROWS:
        row = _add_row(table)
        _label_cell(row.cells[0], label, info_style=True)
        _value_cell(row.cells[1], data.get(key, ""))

    # --- I. OBJECTIVES ---
    _section_header(table, L.HEADER_OBJECTIVES)
    _field_row(table, "A. Content Standard", data.get("content_standard", ""))
    _field_row(table, "B. Performance Standard", data.get("performance_standard", ""))
    _field_row(table, "C. Learning Competencies", data.get("le_competencies_text", ""))

    # --- II. REFERENCES and MATERIALS ---
    _section_header(table, L.HEADER_REFERENCES)
    _field_row(table, "References", data.get("references", ""))
    _field_row(table, "Materials", data.get("learning_resources", ""))

    # --- III. CONTENT ---
    _section_header(table, L.HEADER_CONTENT)
    _merged_value(table, data.get("le_content", ""))

    # --- IV. OBJECTIVES ---
    _section_header(table, L.HEADER_OBJECTIVES2)
    _merged_value(table, _objectives_text(data))

    # --- V. PROCEDURES (own table: wide content | narrow annotations) ---
    _add_spacer(doc)
    ptable = _new_table()
    _fit_grid(ptable, VALUE_W, LABEL_W)
    _add_procedures(ptable, data)

    # --- VI-VIII (main layout continues) ---
    _add_spacer(doc)
    table2 = _new_table()
    _fit_grid(table2, LABEL_W, VALUE_W)

    # --- VI. ASSESSMENT ---
    _section_header(table2, L.HEADER_ASSESSMENT)
    _field_row(table2, "A. Paper and Pen", data.get("le_quiz", ""))
    _field_row(table2, "B. Performance Task", data.get("le_perf_overview", ""))
    _field_row(table2, "Directions to the Learners", data.get("le_perf_directions", ""))
    _field_row(table2, "Scoring Rubrics", data.get("le_perf_rubric", ""))

    # --- VII. REFLECTION ---
    _section_header(table2, L.HEADER_REFLECTION)
    _merged_value(table2, L.REFLECTION_DIRECTIONS, italic=True)
    _field_row(table2, "Reflection", data.get("reflection", ""))

    # --- VIII. USE OF GENERATIVE AI ---
    _section_header(table2, L.HEADER_GENAI)
    _merged_value(table2, data.get("le_ai_declaration", ""))

    # --- Signature block ---
    _add_signature(doc, data)

    doc.save(output_path)
    return output_path
