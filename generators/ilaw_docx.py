"""DOCX export matching the official DepEd ILAW lesson plan template
(single 2-column table, portrait, 8.5in x 13in long bond paper)."""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from generators import ilaw_layout as L
from generators.watermark import add_docx_watermark

LABEL_W = Inches(2.15)
VALUE_W = Inches(4.5)
GRAY = "BFBFBF"


def _set_cell_shading(cell, color):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def _set_cell_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:color"), "000000")
        element.set(qn("w:space"), "0")
        tcBorders.append(element)
    tcPr.append(tcBorders)


def _set_cell_width(cell, width):
    cell.width = width
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = tcPr.find(qn("w:tcW"))
    if tcW is None:
        tcW = OxmlElement("w:tcW")
        tcPr.append(tcW)
    tcW.set(qn("w:w"), str(int(width.twips)))
    tcW.set(qn("w:type"), "dxa")


def _fit_grid(table, width_a, width_b):
    """Set the table grid so Word honors the intended column widths.

    python-docx creates a tblGrid with equal columns; with fixed layout Word
    uses those grid columns instead of the per-cell w:tcW, so without this the
    columns render 50/50 and the layout breaks. We fix gridCol + tblW here.
    """
    table.columns[0].width = width_a
    table.columns[1].width = width_b
    tblPr = table._tbl.tblPr
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), str(int((width_a + width_b) / 635)))
    tblW.set(qn("w:type"), "dxa")


def _para(cell, first=True):
    if first and cell.paragraphs and not cell.paragraphs[0].runs:
        p = cell.paragraphs[0]
    else:
        p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = 1.0
    return p


def _add_runs(p, text, size=10, bold=False, italic=False, underline=False):
    lines = (text or "").split("\n")
    for i, line in enumerate(lines):
        if i > 0:
            run = p.add_run()
            run.add_break()
        run = p.add_run(line)
        run.font.size = Pt(size)
        run.font.name = "Calibri"
        run.bold = bold
        run.italic = italic
        run.underline = underline
    return p


def _label_cell(cell, title, guidance, info_style=False):
    """Fill a left-column label cell.

    info_style=True  -> bold italic label (top info rows)
    info_style=False -> italic underlined field title + small italic guidance
    """
    p = _para(cell)
    if info_style:
        _add_runs(p, title, size=10, bold=True, italic=True)
    else:
        _add_runs(p, title, size=10, italic=True, underline=True)
    if guidance:
        gp = _para(cell, first=False)
        _add_runs(gp, guidance, size=7.5, italic=True)


def _value_cell(cell, value, placeholder="\u2014"):
    p = _para(cell)
    _add_runs(p, value if value else placeholder, size=10)


def _add_row(table, widths=True):
    row = table.add_row()
    for cell in row.cells:
        _set_cell_border(cell)
    if widths:
        _set_cell_width(row.cells[0], LABEL_W)
        _set_cell_width(row.cells[1], VALUE_W)
    return row


def _header_row(table, text):
    row = _add_row(table)
    cell = row.cells[0].merge(row.cells[1])
    _set_cell_shading(cell, GRAY)
    lines = text.split("\n")
    p = _para(cell)
    # First sentence label bold (up to first ':' or '.'), rest regular
    first = lines[0]
    for sep in (":", "."):
        idx = first.find(sep)
        if idx != -1:
            head, tail = first[: idx + 1], first[idx + 1:]
            break
    else:
        head, tail = first, ""
    _add_runs(p, head, size=9.5, bold=True)
    if tail:
        _add_runs(p, tail, size=9.5)
    for line in lines[1:]:
        lp = _para(cell, first=False)
        _add_runs(lp, line, size=9, italic=True)


def _field_row(table, title, guidance, value, placeholder="\u2014"):
    row = _add_row(table)
    _label_cell(row.cells[0], title, guidance)
    _value_cell(row.cells[1], value, placeholder)


def _add_letterhead(doc, data):
    region = data.get("region") or "Region VII - Central Visayas"
    division = data.get("letterhead_division") or ""
    school = data.get("letterhead_school") or ""

    spec = [("Republic of the Philippines", 12, True, False),
            ("Department of Education", 13, True, False),
            (region, 10.5, False, True)]
    if division:
        spec.append((division, 11, True, False))
    if school:
        spec.append((school, 11, True, False))

    first = True
    for text, size, bold, italic in spec:
        p = doc.add_paragraph()
        first = False
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        _add_runs(p, text, size=size, bold=bold, italic=italic)


def _add_rule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom)
    pPr.append(pBdr)


def generate_ilaw_docx(data, output_path, watermark=False):
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.line_spacing = 1.0

    # Portrait, 8.5in x 13in (Philippine long bond), margins per official template
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(13)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(0.8125)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    _add_letterhead(doc, data)
    _add_rule(doc)

    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    _fit_grid(table, LABEL_W, VALUE_W)

    # --- Info rows ---
    for label, guidance, key in L.INFO_ROWS:
        row = _add_row(table)
        _label_cell(row.cells[0], label, guidance, info_style=True)
        _value_cell(row.cells[1], data.get(key, ""))

    # --- Intentions ---
    _header_row(table, L.HEADER_INTENTIONS)
    _field_row(table, "Learning Competency:", L.GUIDE_COMPETENCY, data.get("competency_full", ""))
    _field_row(table, "Learning Objectives:", L.GUIDE_OBJECTIVES, data.get("objectives", ""))
    _field_row(table, "Learner Context:", L.GUIDE_LEARNER_CONTEXT, data.get("learner_context", ""))

    # --- Learning Experience ---
    _header_row(table, L.HEADER_EXPERIENCE)
    _field_row(table, "Pre-lesson:", L.GUIDE_PRE_LESSON, data.get("pre_lesson", ""))

    # Flow row: bold ILAW sub-headings inside the value cell
    row = _add_row(table)
    _label_cell(row.cells[0], "Flow:", L.GUIDE_FLOW)
    vcell = row.cells[1]
    first = True
    for heading, key, time_key in L.FLOW_PARTS:
        hp = _para(vcell, first=first)
        first = False
        _add_runs(hp, L.flow_heading(heading, time_key, data), size=10, bold=True)
        vp = _para(vcell, first=False)
        _add_runs(vp, data.get(key, "") or "\u2014", size=10)

    _field_row(table, "Learning Resources:", L.GUIDE_RESOURCES, data.get("learning_resources", ""))
    _field_row(table, "Opportunities for integration:", L.GUIDE_INTEGRATION,
               data.get("integration", ""), placeholder="N/A")

    # --- Assessment ---
    _header_row(table, L.HEADER_ASSESSMENT)
    _field_row(table, "Formative Assessment:", L.GUIDE_FORMATIVE, data.get("formative_assessment", ""))

    # --- Ways Forward ---
    _header_row(table, L.HEADER_WAYS_FORWARD)
    _field_row(table, "Extended learning opportunities:", L.GUIDE_EXTENDED, data.get("extended_learning", ""))
    _field_row(table, "Reflections:", L.GUIDE_REFLECTIONS, data.get("reflection", ""))

    # --- Signature block ---
    doc.add_paragraph()
    sig = doc.add_table(rows=3, cols=2)
    sig.alignment = WD_TABLE_ALIGNMENT.CENTER

    designation = data.get("designation", "") or "Teacher"
    head_designation = data.get("school_head_designation", "") or "School Head / Principal"
    teacher_name = data.get("teacher", "") or "\u00a0"
    head_name = data.get("school_head", "") or "\u00a0"

    sig_rows = [
        (("Prepared by:", True, False), ("Reviewed & Checked by:", True, False)),
        ((teacher_name, False, True), (head_name, False, True)),
        ((designation, False, False), (head_designation, False, False)),
    ]
    for r, cells in enumerate(sig_rows):
        for c, (text, bold, underline) in enumerate(cells):
            cell = sig.rows[r].cells[c]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if r == 1:
                p.paragraph_format.space_before = Pt(6)
            run = p.add_run(text)
            run.font.size = Pt(10)
            run.font.name = "Calibri"
            run.bold = bold
            run.underline = underline

    if watermark:
        add_docx_watermark(doc)

    doc.save(output_path)
    return output_path
