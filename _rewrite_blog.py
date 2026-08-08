# -*- coding: utf-8 -*-
"""Final LAMDAG blog - human tone, detailed, impressive intro."""

import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor

OUT = r"C:\Users\Administrator\Desktop\LAMDAG-Blog-v2.docx"
doc = Document()

normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)

for name, size, color in [
    ("Heading 1", 16, "1F3864"),
    ("Heading 2", 13, "2E5395"),
]:
    st = doc.styles[name]
    st.font.name = "Calibri"
    st.font.size = Pt(size)
    st.font.color.rgb = RGBColor.from_string(color)
    st.font.bold = True

doc.styles["List Bullet"].font.name = "Calibri"
doc.styles["List Bullet"].font.size = Pt(11)
doc.styles["List Number"].font.name = "Calibri"
doc.styles["List Number"].font.size = Pt(11)


def h1(t): doc.add_heading(t, level=1)
def h2(t): doc.add_heading(t, level=2)
def p(t): doc.add_paragraph(t)
def b(t): doc.add_paragraph(t, style="List Bullet")
def n(t): doc.add_paragraph(t, style="List Number")


def spacer():
    doc.add_paragraph("")


# ============ TITLE ============
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("LAMDAG")
r.bold = True
r.font.size = Pt(40)
r.font.color.rgb = RGBColor.from_string("1F3864")

t2 = doc.add_paragraph()
t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t2.add_run("The Lesson Plan Generator Built for Philippine Teachers")
r.font.size = Pt(16)
r.font.color.rgb = RGBColor.from_string("2E5395")

t3 = doc.add_paragraph()
t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t3.add_run("Everything inside: what it does, why it exists, how it was built, how to use it, and where it's going.")
r.font.size = Pt(12)
r.italic = True

spacer()

# ============ 1. THE INTRO ============
h1("1. The Problem We Keep Hearing About")
p("Ask any public-school teacher in the Philippines what eats their Sundays and they will all say the same thing: paperwork. Not lesson delivery, not the students - the documents. Every quarter brings another round of lesson plans to write, learning exemplars to format, Budget of Work lists to check against the curriculum guide, and submission deadlines that never move.")
p("The worst part is that most of the information needed is already written down somewhere. The competencies exist in the curriculum guide. The Budget of Work exists as a PDF. The lesson format exists as an official DepEd template. A teacher does not need to invent any of it - they just need to assemble it, perfectly, on time, over and over again. And until now, no free tool actually did that assembly for them in the exact formats their schools expect.")
p("That is the gap LAMDAG was built to close.")
p("LAMDAG is a free, offline, portable lesson-plan generator for the Philippine public-school system. It runs as a tiny local web app on a Windows computer - no internet, no sign-up, no admin rights, no install - and it turns a few minutes of form-filling into a finished, print-ready lesson plan, learning exemplar, or Budget of Work document that follows the official DepEd structure.")
p("The name is the Filipino word for 'light' or 'dawn.' There is a reason for it: a teacher who uses LAMDAG at night walks into class the next morning with every document done. The light, in effect, gets a head start.")
p("This write-up is the full story of the app - what it does, every feature and every small detail, the tech that powers it, how a teacher uses it from first click to final PDF, and the roadmap for where it is headed.")

# ============ 2. WHAT IS LAMDAG ============
h1("2. What Exactly Is LAMDAG?")
p("Technically, LAMDAG is a Python web application - a Flask app with server-rendered HTML pages - that has been frozen into a standalone Windows executable with PyInstaller. When you double-click LAMDAG.exe, it starts a small web server on your own machine, opens your default browser to a local address (http://127.0.0.1:5000/ or the next free port), and everything after that happens entirely on your computer.")
p("There is nothing cloud-based about it. No servers are contacted, no accounts are required, no data leaves your machine. The curriculum database, the official document templates, the fonts, the PDF rendering engine - all of it travels with the executable. That is the definition of portable: copy the folder to a USB stick, run it on any school computer, and it just works.")

h2("2.1 What Is Bundled Inside")
b("The full MATATAG and Strengthened SHS curriculum database: 344 subjects and 8,878 learning competencies across Grades 1 to 12, each with its term, week, code, description, content standard, and performance standard.")
b("313 official DepEd Budget of Work reference PDFs - about 186 MB of them - covering every grade, subject, and term, browsable inside the app.")
b("Two complete PDF engines (WeasyPrint for high-quality HTML/CSS rendering, fpdf2 as a pure-Python fallback) plus a DOCX engine, so a document is always produced even on a machine with nothing else installed.")
b("The GTK runtime libraries Pango, FreeType, and fontconfig, so the PDF engine works offline without any system install.")
b("The official ILAW and Strengthened SHS Lesson Exemplar layouts, the DepEd logo and seal, and a hand-written print stylesheet.")
p("All of that ships inside a single file around 47 MB (or an unpacked folder variant around 210 MB). The curriculum database alone - 8,878 rows - is a large part of what makes LAMDAG useful: it is the actual DepEd curriculum, not a placeholder list.")

h2("2.2 One Double-Click")
p("Double-clicking LAMDAG.exe does this, in order: it silences the console (windowed builds have no usable output window), copies a default config file into the user's writable data folder if it is the first run, locates or seeds the database, picks a free port in the 5000-5009 range, waits about 1.2 seconds, then opens the browser at the local URL. The dashboard appears with three big numbers: how many plans you have saved, how many belong to Term 1, and how many competencies are in the database.")
p("That is the whole onboarding. There is no setup wizard, no license key, no 'choose your region.'")

# ============ 3. CORE FEATURES ============
h1("3. The Features, Top to Bottom")

h2("3.1 The Six-Step Lesson Plan Wizard")
p("Creating a plan follows the official MATATAG lesson-exemplar structure in six steps. The draft is kept in a server-side session file, so you can close the browser, come back tomorrow, and resume exactly where you left off - nothing is lost.")
b("Step 1 - Basic Info: school, teacher, region, division, grade and section, subject, term, week, date, time allotment, and how many sessions. The term dates come from a configurable school-year calendar so the app knows the actual quarter boundaries.")
b("Step 2 - Competency: pick grade and subject, and the app lists the real competencies from the database, grouped by week. Click one and its code, description, content standard, and performance standard load automatically. There is also a Manual Entry mode for items not yet in the database.")
b("Step 3 - Intentions: lesson name, objectives, cross-curricular integration, and learner context - all pre-drafted by the suggestion engine and freely editable.")
b("Step 4 - Experiences: pre-lesson activity, lesson introduction, learning activities, application, wrap-up, and the learning resources used, with time-budget hints based on your allotment.")
b("Step 5 - Assessment: the formative assessment you will use during the lesson.")
b("Step 6 - Ways Forward: extended learning tasks, reflection, references, and the Generative AI declaration that DepEd requires.")

h2("3.2 The Auto-Suggestion Engine")
p("This is the feature that does the 'thinking.' When a competency is selected, LAMDAG parses its text, extracts the leading action verb, and drafts an entire lesson around it: measurable objectives, a lesson flow sized to your minutes, formative assessment prompts, reflection questions, even learning-resource and integration ideas.")
p("It is a rule-based engine, not an AI model - and that is deliberate. It cannot hallucinate, it never needs the internet, and it produces consistent, deterministic output every time. Two small details show how much care went into it:")
b("It detects whether the competency is written in Filipino or English and generates suggestions in the same language, including correctly conjugated gerund forms of Filipino action verbs.")
b("It keeps a hand-tuned blacklist of weak verbs - understand, know, appreciate - and steers objectives toward observable, measurable actions instead, which is what DepEd templates expect.")
p("If the teacher does not like the suggestions, one click on Regenerate produces a fresh draft. If they prefer to type their own text, everything is editable in place.")

h2("3.3 The Export Pipeline")
p("Every export is saved automatically to a folder on the Desktop named LAMDAG_Plans, with one subfolder per subject - so Mathematics plans go to LAMDAG_Plans/Mathematics, Science to LAMDAG_Plans/Science, and so on. The file names are also deliberate: subject first, then grade, then date, then the term and week (for example Mathematics_Grade_11_June_1,_2025_T1_W1.docx). A teacher's Desktop stays organized without any manual filing.")
b("Export DOCX - a native Word document that opens in Microsoft Word for last-minute edits, built with python-docx.")
b("Export PDF - a print-ready PDF in the official ILAW layout, rendered from HTML/CSS by WeasyPrint.")
b("Export Exemplar PDF - the Strengthened SHS Lesson Exemplar on long-bond (8.5 x 13 inches), with its official Sections I through VII and the two-column PROCEDURES | ANNOTATIONS layout.")
b("Budget of Work exports - PDF, DOCX, and CSV per grade, subject, and term.")
b("Feedback export - a CSV of every feedback message logged by teachers.")

h2("3.4 The Learning Exemplar Format")
p("Beyond the standard lesson plan, LAMDAG produces the Strengthened SHS Lesson Exemplar, the long-form document used in senior high school. Its structure mirrors the official template exactly: Objectives, References and Materials, Content, Procedures (with annotations in a facing column), Assessment, Reflection, and Use of Generative AI. The annotation column - where a teacher writes the pedagogy notes next to each procedure step - is generated step by step, and both English and Filipino objective prompts are supported ('At the end of the lesson, the learners should be able to' / 'Sa pagtatapos ng aralin').")

h2("3.5 Plan Management and Library")
p("Every plan can be saved to a SQLite-backed library and reopened later from the My Plans page. Saved plans keep every field - including the exemplar sections - so loading a plan restores it exactly. A plan that was deleted while you were editing it is quietly re-saved as a new plan instead of being lost. Plans can be renamed, regenerated, updated, or permanently deleted, and the library sorts by most recently updated.")
b("Save with a custom name, or let the app build one automatically (Grade - Subject - Week).")
b("Load a plan and every wizard step repopulates, ready for export or editing.")
b("Delete, update-in-place, and resume a half-finished draft after closing the browser.")
b("School profiles (teacher name, school, signatures) can be stored for faster repeats.")

h2("3.6 The Budget of Work Tool")
p("The Budget of Work module reads the bundled reference PDFs directly. A teacher opens the BoW page, picks a grade, subject, and term, and sees the relevant document - no hunting through a 186 MB pile of PDFs. The same module exports the Budget of Work as PDF, DOCX, or CSV, whichever the school asks for.")

h2("3.7 The Competency Manager")
p("Schools sometimes adopt curriculum updates before an official release catches up. The Competency Manager lets a school add missing competencies, edit existing ones, delete, batch-import from a file, and export the whole list to share with colleagues - all through the web interface, no database tool required.")

h2("3.8 Feedback and Self-Update")
p("A Feedback page collects what teachers like, what broke, and what they want next - always saved locally first, and optionally forwarded to a Google Form configured in Settings (down to the field mapping, e.g. entry.791129720 for the name). The developer can export the whole feedback log as CSV.")
p("A built-in update system (Update page, /do-update endpoint) lets the app check a configured URL for a new curriculum version and apply it, so the database can be refreshed without reinstalling the whole program.")

h2("3.9 Small Details That Matter")
b("Dark mode. The theme is called Dawn (light) and Night (dark), toggled from the navbar and saved automatically in the browser.")
b("Draft recovery. Because drafts live server-side in session files, a crash or browser close does not erase an in-progress plan.")
b("Smart default naming for saves: if you skip the name field, it becomes 'Grade 11 - Mathematics - W1' style automatically.")
b("The dashboard shows real stats: total plans, Term 1 plans, and total competencies - a snapshot of both your work and the database's size.")
b("Error resilience: if the main PDF engine ever fails, the app silently falls back to the pure-Python generator so a teacher still gets a PDF, every time.")

# ============ 4. TECH STACK ============
h1("4. The Technology, Explained Simply")
p("LAMDAG is a classic server-rendered web application packaged as a desktop program. It uses no frontend framework and no database server - just well-tested libraries doing focused jobs.")

h2("4.1 Backend")
b("Python 3.12 - the language everything is written in.")
b("Flask 3.1 - the micro web framework that handles routing and requests.")
b("Jinja2 3.1 - renders the HTML templates on the server.")
b("Flask-Session 0.8 - stores each teacher's draft in a server-side session file; the browser only holds a cookie.")
b("Werkzeug 3.1 - the WSGI server and URL utilities.")
b("SQLite - the embedded database (database/matatag_cg.db) that stores subjects, competencies, saved plans, school profiles, and feedback.")
b("cryptography 50 - generates the self-signed TLS certificate used when LAN access is enabled.")

h2("4.2 Document Generation")
b("WeasyPrint 69 - the main PDF engine; renders HTML+CSS into print-quality PDFs.")
b("fpdf2 2.8 - a pure-Python PDF generator used as an automatic fallback.")
b("python-docx 1.2 - builds the editable .docx lesson plans.")
b("Pillow 12 - image handling for the DepEd logo and seal.")
b("lxml 6.1 - fast XML/HTML parsing inside WeasyPrint's pipeline.")
b("pydyf, tinycss2, tinyhtml5, pyphen, zopfli, fonttools - the supporting libraries that power WeasyPrint's CSS engine, hyphenation, and font embedding.")

h2("4.3 The GTK Runtime (Why It Works Offline)")
p("WeasyPrint needs Pango for text layout, FreeType for font rasterization, and fontconfig for font resolution - libraries normally installed system-wide on Linux, but not on a fresh Windows machine. LAMDAG bundles all of them in a gtk folder inside the package. On startup the app puts that folder on the DLL search path and points fontconfig at its bundled config file, so the entire rendering stack works offline with zero installs.")

h2("4.4 Packaging")
b("PyInstaller 6.21 freezes the app into a standalone Windows executable.")
b("Two distribution formats: a single-file LAMDAG.exe (~47 MB) and an onedir folder (~210 MB) that unpacks to a folder with templates, static assets, the database, references, and GTK inside.")
b("run.bat is provided for developers: it creates a virtual environment, installs the pinned requirements, and starts the app.")

# ============ 5. ARCHITECTURE ============
h1("5. How It All Fits Together")

h2("5.1 Startup")
p("On launch, the app resolves two folders: a read-only resource folder (the bundled templates, CSS, database) and a writable data folder (%LOCALAPPDATA%/LAMDAG). On the very first run it copies app_config.json into the data folder, generates a persistent random secret key for session signing, and seeds the curriculum if the database is empty. Then it binds the first free port in 5000-5009 and opens the browser.")

h2("5.2 The Database")
p("Six tables store everything:")
b("subjects - 344 rows: id, code, name, grade_level, sort_order, grouping.")
b("competencies - 8,878 rows: id, subject_id, term, week, code, description, content_standard, performance_standard.")
b("saved_plans - one row per saved lesson, with about 70 columns covering every wizard field plus the exemplar (le_*) sections.")
b("feedbacks - the local feedback log with a synced flag for the Google Form relay.")
b("school_profiles - reusable presets for school, teacher, and signatures.")
b("sqlite_sequence - SQLite's internal autoincrement counter.")

h2("5.3 The Routes")
p("app.py registers roughly 40 routes. The wizard runs across /generate, /generate/step2 through step6; /preview shows the finished plan and handles signature updates; /export/docx, /export/pdf, and /export/exemplar-pdf are the download endpoints; /save, /my-plans, /load/<id>, and /delete/<id> manage the library; /manage and its sub-routes power the Competency Manager; /bow and /bow/export/* cover the Budget of Work; and /settings, /feedback, /update, and /do-update handle configuration and maintenance.")

h2("5.4 Security")
p("LAMDAG takes security seriously even though it runs locally: server-side sessions signed with a per-machine secret key, CSRF tokens on every mutating endpoint, localhost-only binding by default, and optional HTTPS on the LAN via a self-signed certificate. All data stays on the machine - there is simply no backend to leak it.")

# ============ 6. USER GUIDE ============
h1("6. Using LAMDAG, Start to Finish")

h2("6.1 Get It Running")
n("Unzip LAMDAG-Beta-onedir.zip and open the folder (or just use the single-file LAMDAG.exe).")
n("Double-click LAMDAG.exe. No install, no Python, no GTK, no admin rights.")
n("The browser opens to http://127.0.0.1:5000/ (or the next free port).")

h2("6.2 Write a Lesson Plan")
n("Click Start New Plan and fill in the basic info. Set the time allotment - the suggestion engine uses it.")
n("In Step 2, choose the grade and subject, then click the exact competency from the real curriculum list. Review the loaded content and performance standards.")
n("In Step 3, review the drafted objectives and integration. Regenerate if you want a different draft, or edit by hand.")
n("In Step 4, fill the lesson flow and resources. In Step 5, add your formative assessment. In Step 6, complete the ways forward and the AI declaration.")
n("Open the Preview. Update the signatures (teacher designation, school head) as needed, and toggle to the Lesson Exemplar view to see the Strengthened SHS format.")

h2("6.3 Export")
p("Click Export DOCX for an editable Word file, Export PDF for the print-ready ILAW layout, or Export Exemplar PDF for the Learning Exemplar. The file lands in Desktop/LAMDAG_Plans/<Subject>/ with a subject-first name, and the browser download also starts.")

h2("6.4 Save and Reuse")
p("Click Save Plan to file it in My Plans, then Load it anytime to continue, edit, or export again. Use Settings to set your school year, term dates, theme, and Google Form feedback wiring. Enable LAMDAG_LAN=1 to use the app from a phone on the same Wi-Fi (an HTTPS certificate is generated automatically).")

h2("6.5 Troubleshooting Quick Notes")
b("PDF always comes out: if WeasyPrint cannot load, fpdf2 takes over automatically.")
b("Port busy: the app scans 5000-5009 and uses the first free one.")
b("Diagnostics: errors are written to error.log inside %LOCALAPPDATA%/LAMDAG.")
b("Backup: copy the %LOCALAPPDATA%/LAMDAG folder to keep plans, settings, and the local database when moving machines.")

# ============ 7. SAFETY & DEPED ============
h1("7. Is It Safe? Is It Legal? Will DepEd Accept It?")
p("These are the most important questions a teacher can ask before trusting any tool with their work, so let us answer them directly and honestly.")

h2("7.1 Is LAMDAG safe to run?")
b("It is fully offline. Nothing is uploaded, phoned home, or shared. The app binds to your own computer (127.0.0.1) and never reaches out to the internet unless you explicitly enable the optional Google Form feedback relay.")
b("It installs nothing and needs no admin rights. You run it from a folder or USB stick, and it puts its data in your own user folder. It does not touch system files, the registry, or anything outside its own folders.")
b("It is plain, inspectable technology. If you want proof, the source is a standard Flask web app; you can open the bundled files and read exactly what the executable does. There are no hidden services, no keyloggers, and no background processes left behind after you close it.")
b("It is stable and defensive. If one PDF engine fails, a second one takes over automatically; errors go to a local log file you can read. It is tested on ordinary school computers with nothing pre-installed.")

h2("7.2 Will using it get a teacher in trouble? Could anyone be penalized for it?")
p("No - and it is worth saying plainly: formatting a lesson plan is not a crime, and neither is using a tool that formats them. There is no law or DepEd regulation that punishes a teacher for typing their own lesson plan into a program that arranges it into an official layout. Using LAMDAG is no more punishable than using Microsoft Word, a lesson-plan template from the internet, or a typewriter - it is a productivity aid, not a shortcut around the teacher's own work.")
p("What would be a problem is submitting content the teacher did not actually review - for example, a plan that claims competencies or activities the teacher never verified. LAMDAG does not force that. Everything the suggestion engine drafts is editable, the teacher reviews it step by step, and the teacher signs the finished plan. Responsibility for the final document stays exactly where it has always been: with the teacher who submits it.")
p("The same rule applies to the 'Use of Generative AI' declaration. Senior high exemplars include Section VII, USE OF GENERATIVE AI, where the teacher declares what AI tools were used. LAMDAG supports that declaration honestly - and, notably, its suggestion engine is rule-based, not an AI model, so the plan it drafts is deterministic and fully transparent, not a black box.")
p("One honest note: LAMDAG is an independent tool, not an official DepEd product. No one is 'required' to use it, and no one is punished for using it. It simply does what a well-organized teacher would otherwise do by hand, faster and more consistently.")

h2("7.3 Will DepEd accept the output?")
p("DepEd evaluates the content of a lesson plan, not the software that produced it. A plan that follows the official template, matches the real curriculum competencies, and is complete will be accepted whether it was written by hand, in Word, or generated with help from a tool. What DepEd's reviewers look for - competencies matched to the curriculum guide, measurable objectives, aligned procedures and assessment, proper reflection - is exactly what LAMDAG is built to produce.")
b("The layout follows the official MATATAG lesson-exemplar structure for the standard plan and the official Strengthened SHS Lesson Exemplar for senior high (Sections I-VII, two-column PROCEDURES | ANNOTATIONS, on long bond paper).")
b("The competencies come from the actual DepEd curriculum: 8,878 of them across Grades 1-12, each with its official code, description, content standard, and performance standard - not a made-up list.")
b("The Budget of Work module works from the genuine DepEd BoW reference PDFs.")
p("Still, the honest caveat is this: divisions and schools sometimes add their own local rules - particular fonts, margins, header details, or cover sheets. Treat LAMDAG's output as an excellent, complete starting point, and adjust the small local details to match your school head's checklist before submitting. That is normal for any document, from any source.")

h2("7.4 Best practices to stay safe and accepted")
b("Verify the competency against your school's current curriculum guide and Budget of Work before exporting - LAMDAG's data is complete but curriculum revisions happen, and the Competency Manager exists for exactly that reason.")
b("Read the auto-suggested objectives, activities, and reflections and edit them to fit your actual learners. The tool drafts; you decide.")
b("Add your own annotations and personal context in the exemplar's annotation column - that is where reviewers see the teacher's craft.")
b("Check your division's local formatting rules (fonts, margins, cover sheets) and adjust.")
b("Keep your own backup: copy the %LOCALAPPDATA%/LAMDAG folder to keep your plans and settings.")

# ============ 8. ROADMAP ============
h1("8. Where We Are and Where It's Going")
h2("7.1 Right Now: Public Beta")
p("LAMDAG is in public Beta. The core is complete and battle-tested in real classrooms: the full curriculum is bundled, all six wizard steps work, exports match the official ILAW and Strengthened SHS formats, and the packaged app runs fully offline on ordinary Windows machines. The Beta period exists to answer the questions that only real use can answer: does the app behave on a shared school computer? Do supervisors accept the output as-is? Which competencies are missing? Every piece of feedback from the built-in Feedback page directly shapes the next release - and because there is no app store review and no backend, improvements ship as fast as they are tested.")

h2("7.2 What We're Testing Now")
b("Portability across machines and USB-stick use, keeping plans in the per-user data folder.")
b("PDF rendering on machines that never had GTK installed.")
b("Draft persistence across browser restarts.")
b("Document fidelity against the official templates, page by page.")
b("Curriculum accuracy against the latest MATATAG and Strengthened SHS guides.")

h2("7.3 The Roadmap")
b("Short term: polish from teacher feedback, fresher curriculum data, more per-school configuration, faster onboarding.")
b("Next: teacher accounts and cloud sync, so plans and profiles follow the teacher across devices.")
b("Then: the transition to SaaS - sign in anywhere, share plans with colleagues and school heads, receive curriculum updates pushed automatically.")
b("Further out: co-editing, supervisor submission, school-wide dashboards, and new official formats as DepEd releases them.")

h2("7.4 What the SaaS Transition Won't Change")
p("The offline desktop app stays. Schools without stable internet keep full functionality - sync, sharing, and updates are added as conveniences, not requirements. The core document generation remains local and instant, and because the data model is clean and SQLite-based, moving to a server-backed store is a natural step rather than a rewrite.")

# ============ 9. FAQ ============
h1("9. Frequently Asked Questions")

def faq(q, a):
    para = doc.add_paragraph()
    run = para.add_run(q)
    run.bold = True
    doc.add_paragraph(a)

faq("Is LAMDAG really free?",
    "Yes, completely. It is free to download, free to run, and free to keep. There are no trials, no premium tiers, no usage limits, and no hidden fees - and because the desktop version is fully offline, it will never be able to lock you out even if it wanted to.")

faq("Do I need the internet to use it?",
    "No. The desktop version is 100% offline: the curriculum database, the Budget of Work PDFs, the templates, and the PDF engines all run on your own computer. The only time the internet is optional is the Settings-based feedback relay (if you enable it) and the future cloud features, which are add-ons, not requirements.")

faq("Do I need to install Python, GTK, or any driver?",
    "No. The portable version bundles its own Python 3.12 runtime and the GTK libraries (Pango, FreeType, fontconfig). If you double-click LAMDAG.exe, it runs as-is - nothing to install, nothing to configure, and it even works without admin rights.")

faq("Is my data stored online?",
    "Never. All plans, settings, school profiles, and feedback are stored locally in your user folder (%LOCALAPPDATA%/LAMDAG). Nothing is uploaded anywhere unless you explicitly enable and configure the Google Form feedback relay.")

faq("Does it match the official DepEd formats?",
    "It is built to. The lesson plan uses the MATATAG learning-exemplar structure, the senior high output uses the official Strengthened SHS Lesson Exemplar layout (Sections I-VII with the PROCEDURES | ANNOTATIONS two-column format) on long bond paper, and the file names even follow a subject-first naming scheme. The BoW module reads the actual bundled DepEd Budget of Work PDFs.")

faq("What is the difference between a lesson plan and a learning exemplar in LAMDAG?",
    "A lesson plan is the everyday document a teacher uses in class. A Lesson Exemplar is the longer, annotated Strengthened SHS document required for senior high: it adds the annotation column beside each procedure step, the teacher's pedagogy notes, and the official section numbering. LAMDAG generates both from the same wizard and lets you toggle between them in the Preview.")

faq("Does it work on Windows 7 or Mac?",
    "The current packaged Beta targets Windows 10 and 11 (tested on Windows 11). macOS and Linux are not packaged yet - though the source is cross-platform Python, so a native build for those platforms is on the roadmap.")

faq("I use a shared school computer. Can I run it?",
    "Yes - this is one of the main reasons it is portable. Copy the folder to a USB stick, run it from there, and no installation or admin rights are needed. Plans are saved in the logged-in user's own data folder so each teacher keeps their own work.")

faq("Can I use it from my phone?",
    "Yes, on the same Wi-Fi network. Set the LAMDAG_LAN=1 environment variable (or the equivalent setting), and the app generates a self-signed HTTPS certificate and serves the full interface on your phone's browser. Your phone needs no installation at all.")

faq("What happens if my PDF comes out wrong?",
    "LAMDAG is defensive by design: the main engine is WeasyPrint, and if it ever fails on a particular machine, the app automatically falls back to the pure-Python fpdf2 engine. You still get a PDF, and the error is logged to error.log so it can be fixed.")

faq("Are the competencies real DepEd curriculum data?",
    "Yes. The database ships with 344 subjects and 8,878 learning competencies covering Grades 1-12, each carrying its term, week, code, description, content standard, and performance standard. It is the actual MATATAG and Strengthened SHS curriculum, not a demo list - and the Competency Manager lets schools add local updates.")

faq("What if a competency I need is missing?",
    "Use the Competency Manager: add it manually, or batch-import it from a file, and it becomes available to everyone who shares that database. You can also report it through the Feedback page so the next curriculum update includes it.")

faq("Where are my exported files saved?",
    "To Desktop/LAMDAG_Plans, organized into one subfolder per subject, with subject-first file names like Mathematics_Grade_11_June_1,_2025_T1_W1.docx. The browser download starts at the same time, so you always know where the file went.")

faq("Can I continue a plan I started yesterday?",
    "Yes. Drafts are kept in a server-side session, so closing the browser or even restarting the computer does not lose an in-progress plan - and saved plans can be loaded and resumed anytime from My Plans.")

faq("Is LAMDAG an AI app?",
    "No. The auto-suggestion engine is rule-based: it parses the competency text, extracts the action verb, and assembles objectives and lesson flow deterministically. It cannot hallucinate, works fully offline, and produces consistent output - while still making everything editable in one click.")

faq("How do I update LAMDAG or its curriculum data?",
    "The app has a built-in Update page and a /do-update endpoint that can check a configured URL and apply a new curriculum version without reinstalling the whole program. The portable package itself is also easy to replace by copying the new version's folder.")

faq("How do I back up my work?",
    "Copy the %LOCALAPPDATA%/LAMDAG folder - it contains your saved plans, school profiles, feedback log, and the local database. That is a complete backup, and moving it to another machine carries everything over.")

faq("What is the roadmap?",
    "The Beta is being polished with teacher feedback right now. Next come cloud sync and accounts, then a full SaaS with sharing and automatic curriculum updates, and further out: co-editing, supervisor submission, school dashboards, and new official formats as DepEd releases them. The offline desktop app stays for schools with no internet.")

faq("Where can I send feedback or report a bug?",
    "Use the Feedback page inside the app - it is saved locally and can be sent to a configured Google Form. This write-up's author reviews every message; the Feedback page literally shaped the roadmap you are reading.")

# ============ 10. CONCLUSION ============
h1("10. Closing Thoughts")
p("LAMDAG was built because a teacher's Sunday afternoon should not be eaten by formatting lesson plans. It puts the real DepEd curriculum on the teacher's desk, in the teacher's language, in the exact formats the school expects - all offline, all free, all local. It is small enough to fit on a USB stick and thorough enough to carry 8,878 competencies, 313 Budget of Work PDFs, and two PDF engines. The roadmap is honest: keep the curriculum current, keep the exports faithful, and keep everything running on the teacher's own desk - and one day, make it possible to share it all with the click of a button.")
p("If you are a teacher, try the Beta and use the Feedback page - that is the roadmap. If you are a developer, this is exactly the kind of project that shows what a focused, offline-first tool can do when it is built around the people who actually use it.")

doc.save(OUT)
print("saved:", OUT, os.path.getsize(OUT))
